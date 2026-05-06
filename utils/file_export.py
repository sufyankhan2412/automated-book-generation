"""
File export utilities — compile chapters into .docx, .pdf, and .txt formats.
"""

from __future__ import annotations
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF


def _safe_path(filepath: Path) -> Path:
    """
    If the file is locked (open in another app), return a new path
    with a timestamp suffix so we don't crash.
    """
    if not filepath.exists():
        return filepath
    # Try writing a test — if PermissionError, use timestamped name
    try:
        with open(filepath, 'ab'):
            pass
        return filepath
    except PermissionError:
        stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        new_path = filepath.with_stem(f'{filepath.stem}_{stamp}')
        print(f'   ⚠️  {filepath.name} is locked — saving as {new_path.name}')
        return new_path


def export_docx(
    title: str,
    chapters: list[dict],
    output_path: Path,
) -> Path:
    """
    Export chapters to a .docx file with proper formatting.
    Handles markdown cleanup and subsection headings.
    """
    import re
    doc = Document()

    # Style setup
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    # Book title page
    doc.add_paragraph()  # spacer
    doc.add_paragraph()  # spacer
    heading = doc.add_heading(title, level=0)
    heading.alignment = 1  # center
    doc.add_page_break()

    for ch in chapters:
        # Chapter heading
        doc.add_heading(
            f"Chapter {ch['chapter_number']}: {ch['chapter_title']}",
            level=1,
        )

        # Clean and format chapter text
        text = _clean_markdown(ch["chapter_text"])

        for paragraph in text.split("\n\n"):
            para = paragraph.strip()
            if not para:
                continue

            # Detect subsection headings (short lines, no period at end, often Title Case)
            lines = para.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Subsection heading: short, no ending period, looks like a title
                if (len(line) < 80 and not line.endswith('.') and not line.endswith(',')
                        and not line.startswith('(') and len(line.split()) >= 2
                        and len(line.split()) <= 12
                        and not any(line.lower().startswith(w) for w in ['the ', 'this ', 'in ', 'a ', 'an ', 'it ', 'as ', 'for ', 'with ', 'but ', 'and ', 'or ', 'to ', 'from '])):
                    doc.add_heading(line, level=2)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    filepath = _safe_path(output_path / f"{_safe_filename(title)}.docx")
    doc.save(str(filepath))
    return filepath


def export_pdf(
    title: str,
    chapters: list[dict],
    output_path: Path,
) -> Path:
    """
    Export chapters to a .pdf file.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title page
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(0, 60, text="", new_x="LMARGIN", new_y="NEXT")  # spacer
    pdf.cell(0, 20, text=title, new_x="LMARGIN", new_y="NEXT", align="C")

    for ch in chapters:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(
            0, 12,
            text=f"Chapter {ch['chapter_number']}: {ch['chapter_title']}",
            new_x="LMARGIN", new_y="NEXT",
        )
        pdf.ln(4)
        pdf.set_font("Helvetica", "", 11)
        # Clean markdown and encode to latin-1 safe text
        clean_text = _clean_markdown(ch["chapter_text"])
        safe_text = clean_text.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 6, text=safe_text)

    filepath = _safe_path(output_path / f"{_safe_filename(title)}.pdf")
    pdf.output(str(filepath))
    return filepath


def export_txt(
    title: str,
    chapters: list[dict],
    output_path: Path,
) -> Path:
    """
    Export chapters to a plain .txt file.
    """
    lines: list[str] = []
    lines.append(f"{'=' * 60}")
    lines.append(f"  {title}")
    lines.append(f"{'=' * 60}\n")

    for ch in chapters:
        lines.append(f"\n{'─' * 50}")
        lines.append(f"  Chapter {ch['chapter_number']}: {ch['chapter_title']}")
        lines.append(f"{'─' * 50}\n")
        lines.append(_clean_markdown(ch["chapter_text"]))
        lines.append("")

    filepath = _safe_path(output_path / f"{_safe_filename(title)}.txt")
    filepath.write_text("\n".join(lines), encoding="utf-8")
    return filepath


def _clean_markdown(text: str) -> str:
    """Strip markdown formatting from LLM output."""
    import re
    # Remove bold: **text** or __text__
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove italic: *text* or _text_
    text = re.sub(r'(?<!\w)\*(.+?)\*(?!\w)', r'\1', text)
    # Remove heading markers: ## or ###
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # Remove horizontal rules: --- or ***
    text = re.sub(r'^[\-\*]{3,}\s*$', '', text, flags=re.MULTILINE)
    # Clean up numbered list bold: **1.** or 1. **text**
    text = re.sub(r'\*\*(\d+\.)\*\*', r'\1', text)
    return text.strip()


def _safe_filename(title: str) -> str:
    """Convert a title into a filesystem-safe name."""
    return "".join(c if c.isalnum() or c in " -_" else "" for c in title).strip().replace(" ", "_")
